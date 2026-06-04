"""DOCX 생성 — Markdown 을 명명 스타일에 매핑(python-docx).

스펙(회사 양식 가이드)과 동일한 매핑:
  # / ## / ###  → 제목 1 / 제목 2 / 제목 3  (없으면 Heading 1~3)
  일반 단락       → 본문 (없으면 Normal)
  불릿           → 목록 (없으면 List Bullet)
  **굵게**        → 런 bold (강조 char 스타일 있으면 적용)
템플릿(.docx)이 있으면 그 스타일/표지를 상속한다.
"""
from __future__ import annotations

import os
from typing import Any, Dict, List, Optional

from .. import mdblocks
from . import md_gen


# 한글 스타일명 → 폴백(영문 기본 스타일)
_HEADING_NAMES = {1: ["제목 1", "Heading 1"], 2: ["제목 2", "Heading 2"], 3: ["제목 3", "Heading 3"],
                  4: ["제목 4", "Heading 4"], 5: ["제목 5", "Heading 5"], 6: ["제목 6", "Heading 6"]}
_BODY_NAMES = ["본문", "Normal"]
_BULLET_NAMES = ["목록", "List Bullet"]
_QUOTE_NAMES = ["인용", "Quote", "Intense Quote"]
_TITLE_NAMES = ["제목", "Title"]


def _style(doc, names: List[str]):
    """문서에 존재하는 첫 스타일명을 반환(없으면 None)."""
    have = {s.name for s in doc.styles}
    for n in names:
        if n in have:
            return n
    return None


def render(payload: Any, out_path: str, template: Optional[str] = None) -> str:
    from docx import Document

    md = md_gen.to_markdown(payload)
    doc = Document(template) if (template and os.path.exists(template)) else Document()

    # 템플릿에 표지 placeholder 가 없을 때, 첫 H1 을 Title 스타일로
    title_style = _style(doc, _TITLE_NAMES)
    body_style = _style(doc, _BODY_NAMES)
    bullet_style = _style(doc, _BULLET_NAMES)
    quote_style = _style(doc, _QUOTE_NAMES)

    blocks: List[Dict] = mdblocks.parse(md)
    used_title = False
    for blk in blocks:
        btype = blk["type"]
        if btype == "heading":
            lvl = min(max(blk["level"], 1), 6)
            if lvl == 1 and not used_title and title_style:
                p = doc.add_paragraph(style=title_style)
                used_title = True
            else:
                hs = _style(doc, _HEADING_NAMES[lvl])
                p = doc.add_paragraph(style=hs) if hs else doc.add_heading(level=lvl)
                if hs is None:
                    p.clear()  # add_heading 이 텍스트 없는 단락 생성 → runs 로 채움
            _add_runs(p, blk["runs"])
        elif btype == "bullet":
            p = doc.add_paragraph(style=bullet_style) if bullet_style else doc.add_paragraph(style=None)
            if not bullet_style:
                p.add_run("• ")
            _add_runs(p, blk["runs"])
        elif btype == "quote":
            p = doc.add_paragraph(style=quote_style) if quote_style else doc.add_paragraph()
            _add_runs(p, blk["runs"])
        elif btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(blk["text"])
            run.font.name = "D2Coding"
        elif btype == "table":
            ncol = blk.get("ncol") or 0
            if ncol:
                table = doc.add_table(rows=0, cols=ncol)
                try:
                    table.style = "Table Grid"
                except Exception:
                    pass
                if blk.get("header"):
                    cells = table.add_row().cells
                    for j, h in enumerate(blk["header"][:ncol]):
                        cells[j].text = h
                        for para in cells[j].paragraphs:
                            for r in para.runs:
                                r.bold = True
                for row in blk.get("rows", []):
                    cells = table.add_row().cells
                    for j, v in enumerate(row[:ncol]):
                        cells[j].text = v
        else:  # para
            p = doc.add_paragraph(style=body_style) if body_style else doc.add_paragraph()
            _add_runs(p, blk["runs"])

    doc.save(out_path)
    return out_path


def _add_runs(paragraph, runs: List) -> None:
    for text, bold in runs:
        if not text:
            continue
        r = paragraph.add_run(text)
        if bold:
            r.bold = True
